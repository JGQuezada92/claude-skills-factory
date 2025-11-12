#!/usr/bin/env python3
"""
Word Report Generator

Creates professional Microsoft Word (.docx) reports for liquidity analysis
with proper formatting, sections, tables, and styling suitable for investment committees.
"""

from typing import Dict, List, Optional
from datetime import datetime
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")


class WordReportGenerator:
    """Generate Word documents for liquidity analysis reports"""
    
    def __init__(self):
        self.document = None
        if not DOCX_AVAILABLE:
            print("Error: python-docx is required for Word generation. Install with: pip install python-docx")
    
    def create_liquidity_report(
        self,
        analysis_results: Dict,
        output_path: str = 'Global_Liquidity_Analysis_Report.docx'
    ) -> str:
        """
        Create comprehensive liquidity analysis Word report
        
        Args:
            analysis_results: Dictionary containing analysis results with sections:
                - executive_summary: Executive summary text
                - liquidity_cycle: Global liquidity cycle assessment
                - central_bank_policy: Central bank policy analysis
                - cross_border_flows: Cross-border capital flows analysis
                - monetary_aggregates: Monetary aggregates review
                - collateral_market: Collateral market health assessment
                - asset_price_implications: Asset price implications analysis
                - forward_looking_forecast: Forward-looking forecast
            output_path: Path to save Word document
            
        Returns:
            Path to created Word file
        """
        if not DOCX_AVAILABLE:
            print("Error: python-docx is required for Word generation.")
            return ""
        
        try:
            self.document = Document()
            
            # Set up document styles
            self._setup_styles()
            
            # Create title page
            self._create_title_page()
            
            # Create Executive Summary
            if 'executive_summary' in analysis_results:
                self._add_section("Executive Summary", analysis_results['executive_summary'])
            
            # Create Global Liquidity Cycle Assessment
            if 'liquidity_cycle' in analysis_results:
                self._add_section("Global Liquidity Cycle Assessment", analysis_results['liquidity_cycle'])
            
            # Create Central Bank Policy Analysis
            if 'central_bank_policy' in analysis_results:
                self._add_section("Central Bank Policy Analysis", analysis_results['central_bank_policy'])
            
            # Create Cross-Border Capital Flows
            if 'cross_border_flows' in analysis_results:
                self._add_section("Cross-Border Capital Flows", analysis_results['cross_border_flows'])
            
            # Create Monetary Aggregates Review
            if 'monetary_aggregates' in analysis_results:
                self._add_section("Monetary Aggregates Review", analysis_results['monetary_aggregates'])
            
            # Create Collateral Market Health
            if 'collateral_market' in analysis_results:
                self._add_section("Collateral Market Health", analysis_results['collateral_market'])
            
            # Create Asset Price Implications
            if 'asset_price_implications' in analysis_results:
                self._add_section("Asset Price Implications", analysis_results['asset_price_implications'])
            
            # Create Forward-Looking Forecast
            if 'forward_looking_forecast' in analysis_results:
                self._add_section("Forward-Looking Forecast", analysis_results['forward_looking_forecast'])
            
            # Add data sources and citations
            self._add_data_sources(analysis_results.get('data_sources', []))
            
            # Save document
            self.document.save(output_path)
            print(f"✓ Word report created: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"Error creating Word report: {e}")
            return ""
    
    def _setup_styles(self):
        """Set up document styles"""
        styles = self.document.styles
        
        # Heading 1 style
        if 'Heading 1' in styles:
            h1 = styles['Heading 1']
            h1.font.name = 'Calibri'
            h1.font.size = Pt(18)
            h1.font.bold = True
            h1.font.color.rgb = RGBColor(54, 96, 146)  # Dark blue
        
        # Heading 2 style
        if 'Heading 2' in styles:
            h2 = styles['Heading 2']
            h2.font.name = 'Calibri'
            h2.font.size = Pt(14)
            h2.font.bold = True
            h2.font.color.rgb = RGBColor(54, 96, 146)
    
    def _create_title_page(self):
        """Create title page for the report"""
        # Title
        title = self.document.add_heading('GLOBAL MARKET LIQUIDITY ANALYSIS', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format = title.runs[0]
        title_format.font.size = Pt(24)
        title_format.font.bold = True
        title_format.font.color.rgb = RGBColor(54, 96, 146)
        
        # Add spacing
        self.document.add_paragraph()
        
        # Date
        date_para = self.document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now().strftime('%B %Y'))
        date_run.font.size = Pt(14)
        date_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Page break
        self.document.add_page_break()
    
    def _add_section(self, section_title: str, content: str):
        """Add a section to the document"""
        # Section heading
        heading = self.document.add_heading(section_title, level=1)
        
        # Section content
        # Split content into paragraphs if it contains newlines
        paragraphs = content.split('\n\n') if content else []
        
        for para_text in paragraphs:
            if para_text.strip():
                para = self.document.add_paragraph(para_text.strip())
                para_format = para.runs[0] if para.runs else None
                if para_format:
                    para_format.font.name = 'Calibri'
                    para_format.font.size = Pt(11)
        
        # Add spacing between sections
        self.document.add_paragraph()
    
    def _add_data_sources(self, data_sources: List[str]):
        """Add data sources section"""
        if not data_sources:
            return
        
        self.document.add_page_break()
        heading = self.document.add_heading('Data Sources and Citations', level=1)
        
        # Add sources as bulleted list
        for source in data_sources:
            para = self.document.add_paragraph(source, style='List Bullet')
            para_format = para.runs[0] if para.runs else None
            if para_format:
                para_format.font.name = 'Calibri'
                para_format.font.size = Pt(10)
                para_format.font.color.rgb = RGBColor(100, 100, 100)
    
    def add_table(self, headers: List[str], rows: List[List[str]], title: Optional[str] = None):
        """Add a table to the document"""
        if title:
            self.document.add_heading(title, level=2)
        
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        
        self.document.add_paragraph()


if __name__ == "__main__":
    # Example usage
    print("Word Report Generator")
    print("=" * 50)
    print("This script creates Word (.docx) documents for liquidity analysis reports")
    print("Import this module and use WordReportGenerator class in your analysis")
    print()
    print("Note: Requires python-docx library")
    print("Install with: pip install python-docx")

