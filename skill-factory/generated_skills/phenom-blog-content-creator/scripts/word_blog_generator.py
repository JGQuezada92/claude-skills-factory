#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Phenom Blog Word Document Generator
Generates professionally formatted Word documents matching Phenom's visual style
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Optional


class PhenomBlogGenerator:
    """Generates Phenom-style blog posts as Word documents"""
    
    # Phenom brand colors (extracted from visual examples)
    PRIMARY_COLOR = RGBColor(0, 51, 102)  # Dark blue
    ACCENT_COLOR = RGBColor(0, 102, 204)  # Medium blue
    TEXT_COLOR = RGBColor(51, 51, 51)  # Dark gray
    
    def __init__(self, output_path: str = "Phenom_Blog_Post.docx"):
        """
        Initialize the blog generator
        
        Args:
            output_path: Path where the Word document will be saved
        """
        self.doc = Document()
        self.output_path = output_path
        self._setup_styles()
    
    def _setup_styles(self):
        """Set up custom styles matching Phenom's visual formatting"""
        styles = self.doc.styles
        
        # Title style
        if 'Phenom Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Phenom Title', 1)  # 1 = paragraph style
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(32)
            title_font.bold = True
            title_font.color.rgb = self.TEXT_COLOR
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_style.paragraph_format.space_after = Pt(12)
        
        # Heading 1 style (main sections)
        if 'Phenom Heading 1' not in [s.name for s in styles]:
            h1_style = styles.add_style('Phenom Heading 1', 2)  # 2 = character style
            h1_font = h1_style.font
            h1_font.name = 'Calibri'
            h1_font.size = Pt(24)
            h1_font.bold = True
            h1_font.color.rgb = self.TEXT_COLOR
        
        # Heading 2 style (subsections)
        if 'Phenom Heading 2' not in [s.name for s in styles]:
            h2_style = styles.add_style('Phenom Heading 2', 2)
            h2_font = h2_style.font
            h2_font.name = 'Calibri'
            h2_font.size = Pt(18)
            h2_font.bold = True
            h2_font.color.rgb = self.TEXT_COLOR
        
        # Body text style
        if 'Phenom Body' not in [s.name for s in styles]:
            body_style = styles.add_style('Phenom Body', 1)
            body_font = body_style.font
            body_font.name = 'Calibri'
            body_font.size = Pt(11)
            body_font.color.rgb = self.TEXT_COLOR
            body_style.paragraph_format.line_spacing = 1.15
            body_style.paragraph_format.space_after = Pt(12)
        
        # Quote style
        if 'Phenom Quote' not in [s.name for s in styles]:
            quote_style = styles.add_style('Phenom Quote', 1)
            quote_font = quote_style.font
            quote_font.name = 'Calibri'
            quote_font.size = Pt(11)
            quote_font.italic = True
            quote_font.color.rgb = self.TEXT_COLOR
            quote_style.paragraph_format.left_indent = Inches(0.5)
            quote_style.paragraph_format.space_after = Pt(12)
    
    def add_title(self, title: str):
        """Add blog post title"""
        paragraph = self.doc.add_paragraph(title)
        paragraph.style = 'Phenom Title'
        return paragraph
    
    def add_heading(self, text: str, level: int = 1):
        """Add section heading"""
        if level == 1:
            paragraph = self.doc.add_heading(text, level=1)
            paragraph.style = 'Phenom Heading 1'
        elif level == 2:
            paragraph = self.doc.add_heading(text, level=2)
            paragraph.style = 'Phenom Heading 2'
        else:
            paragraph = self.doc.add_heading(text, level=level)
        return paragraph
    
    def add_paragraph(self, text: str, style: str = 'Phenom Body'):
        """Add body paragraph"""
        paragraph = self.doc.add_paragraph(text, style=style)
        return paragraph
    
    def add_quote(self, quote_text: str, attribution: Optional[str] = None):
        """Add quoted text with optional attribution"""
        paragraph = self.doc.add_paragraph(f'"{quote_text}"', style='Phenom Quote')
        if attribution:
            attribution_para = self.doc.add_paragraph(f'— {attribution}', style='Phenom Body')
            attribution_para.paragraph_format.left_indent = Inches(0.5)
        return paragraph
    
    def add_bullet_list(self, items: List[str], bold_items: Optional[List[str]] = None):
        """Add bulleted list with optional bold items"""
        if bold_items is None:
            bold_items = []
        
        for item in items:
            paragraph = self.doc.add_paragraph(style='Phenom Body')
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            
            # Check if item should have bold prefix
            item_text = item
            for bold_term in bold_items:
                if item.startswith(bold_term):
                    item_text = item
                    parts = item.split(':', 1)
                    if len(parts) == 2:
                        run = paragraph.add_run(parts[0] + ':')
                        run.bold = True
                        paragraph.add_run(parts[1])
                    else:
                        paragraph.add_run(item)
                    break
            else:
                paragraph.add_run(item)
    
    def add_numbered_list(self, items: List[str]):
        """Add numbered list"""
        for i, item in enumerate(items, 1):
            paragraph = self.doc.add_paragraph(f'{i}. {item}', style='Phenom Body')
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    
    def add_link(self, text: str, url: str, bold: bool = False):
        """Add hyperlink"""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.hyperlink = url
        if bold:
            run.bold = True
        run.font.color.rgb = self.ACCENT_COLOR
        return paragraph
    
    def add_divider(self):
        """Add visual divider (* * *)"""
        paragraph = self.doc.add_paragraph('* * *')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
        return paragraph
    
    def add_table_of_contents(self, items: List[str]):
        """Add table of contents section"""
        self.add_heading('In This Article', level=2)
        for item in items:
            paragraph = self.doc.add_paragraph(f'- {item}', style='Phenom Body')
            paragraph.paragraph_format.left_indent = Inches(0.25)
    
    def save(self):
        """Save the document"""
        self.doc.save(self.output_path)
        return self.output_path


def generate_phenom_blog_word(
    title: str,
    introduction: str,
    table_of_contents: List[str],
    sections: List[Dict],
    conclusion: str,
    output_path: str = "Phenom_Blog_Post.docx"
) -> str:
    """
    Generate a complete Phenom-style blog post as a Word document
    
    Args:
        title: Blog post title
        introduction: Introduction paragraph(s)
        table_of_contents: List of section names for TOC
        sections: List of section dictionaries with structure:
            {
                'heading': 'Section Title',
                'content': 'Section content...',
                'quotes': [{'text': 'quote', 'attribution': 'Name'}],
                'lists': [{'type': 'bullet', 'items': [...]}],
                'links': [{'text': 'Link Text', 'url': 'https://...'}]
            }
        conclusion: Conclusion paragraph(s)
        output_path: Path to save the Word document
    
    Returns:
        Path to saved document
    """
    generator = PhenomBlogGenerator(output_path)
    
    # Add title
    generator.add_title(title)
    
    # Add introduction
    generator.add_paragraph(introduction)
    
    # Add table of contents
    if table_of_contents:
        generator.add_table_of_contents(table_of_contents)
    
    # Add sections
    for section in sections:
        generator.add_heading(section.get('heading', ''), level=1)
        
        # Add content paragraphs
        if 'content' in section:
            paragraphs = section['content'].split('\n\n')
            for para in paragraphs:
                if para.strip():
                    generator.add_paragraph(para.strip())
        
        # Add quotes
        if 'quotes' in section:
            for quote in section['quotes']:
                generator.add_quote(
                    quote.get('text', ''),
                    quote.get('attribution')
                )
        
        # Add lists
        if 'lists' in section:
            for list_item in section['lists']:
                if list_item.get('type') == 'bullet':
                    generator.add_bullet_list(
                        list_item.get('items', []),
                        list_item.get('bold_items', [])
                    )
                elif list_item.get('type') == 'numbered':
                    generator.add_numbered_list(list_item.get('items', []))
        
        # Add links
        if 'links' in section:
            for link in section['links']:
                generator.add_link(
                    link.get('text', ''),
                    link.get('url', ''),
                    link.get('bold', False)
                )
        
        # Add divider between sections (except last)
        if section != sections[-1]:
            generator.add_divider()
    
    # Add conclusion
    if conclusion:
        generator.add_divider()
        generator.add_paragraph(conclusion)
    
    # Save document
    return generator.save()


if __name__ == '__main__':
    # Example usage
    sections = [
        {
            'heading': 'Example Company',
            'content': 'This is an example section with content about the company and their challenges.',
            'quotes': [
                {'text': 'Our biggest challenge is...', 'attribution': 'John Doe, Director'}
            ],
            'lists': [
                {
                    'type': 'bullet',
                    'items': [
                        '295% increase in applications',
                        '800 hours saved',
                        'Improved engagement'
                    ]
                }
            ],
            'links': [
                {'text': 'Learn more', 'url': 'https://example.com', 'bold': True}
            ]
        }
    ]
    
    generate_phenom_blog_word(
        title='Example Blog Post Title',
        introduction='This is an introduction paragraph that sets the stage.',
        table_of_contents=['Example Company'],
        sections=sections,
        conclusion='This is a conclusion paragraph.',
        output_path='example_phenom_blog.docx'
    )


