#!/usr/bin/env python3
"""
Word Report Generator for HCM Industry Analysis

Creates professional Microsoft Word (.docx) reports with Arial font, wide margins,
structured layout including table of contents, clearly outlined headers, bullet points,
and bold formatting for high-value strategic insights suitable for executive consumption.
"""

from typing import Dict, List, Optional, Any
from datetime import datetime
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")


class WordReportGenerator:
    """Generate Word documents for HCM industry analysis reports"""
    
    def __init__(self):
        self.document = None
        self.sections_for_toc = []
        if not DOCX_AVAILABLE:
            print("Error: python-docx is required for Word generation. Install with: pip install python-docx")
    
    def create_hcm_analysis_report(
        self,
        title: str = "HCM Industry Analysis Report",
        author_name: str = "HCM Industry Expert",
        author_title: str = "Investment Banker & Management Consultant",
        sections: List[Dict] = None,
        output_path: str = 'HCM_Industry_Analysis_Report.docx'
    ) -> str:
        """
        Create comprehensive HCM analysis Word report
        
        Args:
            title: Main title for the report
            author_name: Author name
            author_title: Author title/role
            sections: List of section dictionaries with structure:
                [
                    {
                        'heading': 'Section Title',
                        'content': 'Section content text...',
                        'key_insights': ['Key insight 1', 'Key insight 2'],  # Will be bolded
                        'recommendations': ['Recommendation 1', 'Recommendation 2'],  # Will be bolded
                        'subsections': [
                            {
                                'heading': 'Subsection Title',
                                'content': 'Subsection content...',
                                'key_points': ['Point 1', 'Point 2']  # Will be bulleted
                            }
                        ]
                    }
                ]
            output_path: Path to save Word document
            
        Returns:
            Path to saved Word file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        self.document = Document()
        self.sections_for_toc = []
        
        # Set up document with wide margins
        self._setup_document_formatting()
        
        # Create title page
        self._create_title_page(title, author_name, author_title)
        
        # Create table of contents placeholder (will be updated)
        toc_placeholder = self._create_table_of_contents_placeholder()
        
        # Add sections
        if sections:
            for section in sections:
                self._add_section(section, level=1)
        
        # Update table of contents with actual sections
        self._update_table_of_contents(toc_placeholder)
        
        # Add data sources section if needed
        self._add_data_sources_section()
        
        # Save document
        self.document.save(output_path)
        print(f"✓ Word report created: {output_path}")
        return output_path
    
    def _setup_document_formatting(self):
        """Set up document with wide margins and Arial font"""
        # Set wide margins (1.5 inches on all sides)
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(1.5)
            section.bottom_margin = Inches(1.5)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(1.5)
        
        # Set default font to Arial
        self._set_default_font('Arial')
    
    def _set_default_font(self, font_name: str):
        """Set default font for the document"""
        styles = self.document.styles
        
        # Set Normal style to Arial
        if 'Normal' in styles:
            normal_style = styles['Normal']
            normal_style.font.name = font_name
            normal_style.font.size = Pt(11)
            normal_style._element.rPr.rFonts.set(qn('w:ascii'), font_name)
            normal_style._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
        
        # Set Heading 1 style
        if 'Heading 1' in styles:
            h1 = styles['Heading 1']
            h1.font.name = font_name
            h1.font.size = Pt(18)
            h1.font.bold = True
            h1._element.rPr.rFonts.set(qn('w:ascii'), font_name)
            h1._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
        
        # Set Heading 2 style
        if 'Heading 2' in styles:
            h2 = styles['Heading 2']
            h2.font.name = font_name
            h2.font.size = Pt(14)
            h2.font.bold = True
            h2._element.rPr.rFonts.set(qn('w:ascii'), font_name)
            h2._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
        
        # Set Heading 3 style
        if 'Heading 3' in styles:
            h3 = styles['Heading 3']
            h3.font.name = font_name
            h3.font.size = Pt(12)
            h3.font.bold = True
            h3._element.rPr.rFonts.set(qn('w:ascii'), font_name)
            h3._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    
    def _create_title_page(self, title: str, author_name: str, author_title: str):
        """Create title page with report title, author, and date"""
        # Title
        title_para = self.document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title.upper())
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add spacing
        self.document.add_paragraph()
        self.document.add_paragraph()
        
        # Author name
        author_para = self.document.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(author_name)
        author_run.font.name = 'Arial'
        author_run.font.size = Pt(14)
        author_run.font.bold = True
        
        # Author title
        title_para2 = self.document.add_paragraph()
        title_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run2 = title_para2.add_run(author_title)
        title_run2.font.name = 'Arial'
        title_run2.font.size = Pt(12)
        
        # Date
        self.document.add_paragraph()
        date_para = self.document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now().strftime('%B %d, %Y'))
        date_run.font.name = 'Arial'
        date_run.font.size = Pt(11)
        date_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Page break
        self.document.add_page_break()
    
    def _create_table_of_contents_placeholder(self) -> int:
        """Create placeholder for table of contents"""
        toc_heading = self.document.add_heading('Table of Contents', level=1)
        toc_heading.style.font.name = 'Arial'
        
        # Add placeholder paragraph
        toc_para = self.document.add_paragraph()
        toc_para.add_run("Table of contents will be generated here...")
        toc_para.runs[0].font.name = 'Arial'
        toc_para.runs[0].font.size = Pt(11)
        toc_para.runs[0].font.italic = True
        toc_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
        
        # Store paragraph index for later replacement
        toc_index = len(self.document.paragraphs) - 1
        
        # Page break after TOC
        self.document.add_page_break()
        
        return toc_index
    
    def _update_table_of_contents(self, toc_placeholder_index: int):
        """Update table of contents with actual sections"""
        # Remove placeholder paragraph
        if toc_placeholder_index < len(self.document.paragraphs):
            toc_para = self.document.paragraphs[toc_placeholder_index]
            # Clear the paragraph
            for run in toc_para.runs:
                toc_para._element.remove(run._element)
            
            # Add TOC entries
            for i, section_info in enumerate(self.sections_for_toc, 1):
                toc_entry = toc_para.add_run(f"{i}. {section_info['heading']}")
                toc_entry.font.name = 'Arial'
                toc_entry.font.size = Pt(11)
                
                # Add page number reference (simplified - actual TOC would need field codes)
                toc_entry.add_tab()
                toc_page = toc_para.add_run(f"Page {section_info.get('page', 'N/A')}")
                toc_page.font.name = 'Arial'
                toc_page.font.size = Pt(11)
                toc_page.font.color.rgb = RGBColor(100, 100, 100)
                
                # Add line break for next entry
                if i < len(self.sections_for_toc):
                    toc_para.add_run('\n')
    
    def _add_section(self, section: Dict, level: int = 1):
        """Add a section to the document with proper formatting"""
        heading_text = section.get('heading', 'Section')
        
        # Add heading
        heading = self.document.add_heading(heading_text, level=level)
        heading.style.font.name = 'Arial'
        
        # Track for TOC (only level 1 headings)
        if level == 1:
            self.sections_for_toc.append({'heading': heading_text, 'page': 1})  # Simplified page tracking
        
        # Add content
        content = section.get('content', '')
        if content:
            paragraphs = content.split('\n\n') if content else []
            for para_text in paragraphs:
                if para_text.strip():
                    para = self.document.add_paragraph(para_text.strip())
                    self._format_paragraph(para, 'Arial', Pt(11))
        
        # Add key insights (bolded)
        key_insights = section.get('key_insights', [])
        if key_insights:
            insights_heading = self.document.add_heading('Key Strategic Insights', level=level + 1)
            insights_heading.style.font.name = 'Arial'
            for insight in key_insights:
                para = self.document.add_paragraph(insight, style='List Bullet')
                # Make the insight bold
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add recommendations (bolded)
        recommendations = section.get('recommendations', [])
        if recommendations:
            rec_heading = self.document.add_heading('Recommendations', level=level + 1)
            rec_heading.style.font.name = 'Arial'
            for rec in recommendations:
                para = self.document.add_paragraph(rec, style='List Bullet')
                # Make the recommendation bold
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add key metrics (bolded)
        key_metrics = section.get('key_metrics', [])
        if key_metrics:
            metrics_heading = self.document.add_heading('Key Metrics', level=level + 1)
            metrics_heading.style.font.name = 'Arial'
            for metric in key_metrics:
                para = self.document.add_paragraph(metric, style='List Bullet')
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.bold = True
        
        # Add subsections
        subsections = section.get('subsections', [])
        for subsection in subsections:
            self._add_section(subsection, level=level + 1)
        
        # Add spacing between sections
        self.document.add_paragraph()
    
    def _format_paragraph(self, paragraph, font_name: str, font_size: Pt):
        """Format a paragraph with specified font"""
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
        # If paragraph has no runs, add one
        if not paragraph.runs:
            run = paragraph.add_run(paragraph.text)
            run.font.name = font_name
            run.font.size = font_size
            paragraph.text = ''  # Clear original text
    
    def _add_data_sources_section(self):
        """Add data sources and citations section"""
        self.document.add_page_break()
        heading = self.document.add_heading('Data Sources and Citations', level=1)
        heading.style.font.name = 'Arial'
        
        # Placeholder for data sources
        para = self.document.add_paragraph("Data sources will be listed here with proper citations.")
        self._format_paragraph(para, 'Arial', Pt(11))
    
    def add_table(self, headers: List[str], rows: List[List[str]], title: Optional[str] = None):
        """Add a table to the document"""
        if title:
            heading = self.document.add_heading(title, level=2)
            heading.style.font.name = 'Arial'
        
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_para = header_cells[i].paragraphs[0]
            header_run = header_para.runs[0]
            header_run.font.name = 'Arial'
            header_run.font.size = Pt(11)
            header_run.font.bold = True
        
        # Add rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                row_para = row_cells[i].paragraphs[0]
                row_run = row_para.runs[0]
                row_run.font.name = 'Arial'
                row_run.font.size = Pt(10)
        
        self.document.add_paragraph()


def create_hcm_analysis_report(
    title: str = "HCM Industry Analysis Report",
    author_name: str = "HCM Industry Expert",
    author_title: str = "Investment Banker & Management Consultant",
    sections: List[Dict] = None,
    output_path: str = 'HCM_Industry_Analysis_Report.docx'
) -> str:
    """
    Convenience function to create HCM analysis Word report
    
    Args:
        title: Main title for the report
        author_name: Author name
        author_title: Author title/role
        sections: List of section dictionaries
        output_path: Path to save Word file
        
    Returns:
        Path to saved Word file
    """
    generator = WordReportGenerator()
    return generator.create_hcm_analysis_report(
        title=title,
        author_name=author_name,
        author_title=author_title,
        sections=sections,
        output_path=output_path
    )


if __name__ == "__main__":
    # Example usage
    generator = WordReportGenerator()
    
    sections = [
        {
            'heading': 'Executive Summary',
            'content': 'This report provides comprehensive analysis of the HCM industry...',
            'key_insights': [
                'HCM market showing strong growth with AI integration trends',
                'Mid-market segment represents significant opportunity',
                'Competitive landscape is fragmenting with new entrants'
            ],
            'recommendations': [
                'Focus on vertical specialization for mid-market',
                'Invest in AI-powered candidate matching capabilities',
                'Develop channel partner strategy for geographic expansion'
            ]
        },
        {
            'heading': 'Market Analysis',
            'content': 'The HCM market is experiencing significant transformation...',
            'key_metrics': [
                'Global HCM market: $25-30B growing at 8-12% CAGR',
                'Talent Acquisition segment: Largest at ~$8-10B',
                'Mid-market win rates: 40-50% vs. 20-30% enterprise'
            ],
            'subsections': [
                {
                    'heading': 'Market Segmentation',
                    'content': 'Market divided into six key categories...'
                }
            ]
        }
    ]
    
    generator.create_hcm_analysis_report(
        title="HCM Industry Analysis Report",
        author_name="HCM Industry Expert",
        author_title="Investment Banker & Management Consultant",
        sections=sections,
        output_path='HCM_Industry_Analysis_Report.docx'
    )


